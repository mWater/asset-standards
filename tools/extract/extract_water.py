# -*- coding: utf-8 -*-
import docx, csv, re, io, json, unicodedata
from collections import OrderedDict, defaultdict

D = docx.Document('source.docx')
T = D.tables
LOG = []
def log(*a): LOG.append(' '.join(str(x) for x in a))

def cell(c):
    # verbatim text; normalise the docx cell's internal newlines to a single space
    return re.sub(r'[ \t]+', ' ', c.text.replace('\n', ' ').replace('\r', ' ')).strip()

def snake(s):
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(ch for ch in s if not unicodedata.combining(ch))
    s = s.lower()
    s = re.sub(r'[^a-z0-9]+', '_', s)
    return s.strip('_')

def kind(row):
    """B=group, N=attribute-or-plain, I=choice, based on run formatting of the name cell."""
    nc = row.cells[1]
    b = i = n = 0
    for p in nc.paragraphs:
        for r in p.runs:
            if not r.text.strip(): continue
            if r.bold: b += 1
            elif r.italic: i += 1
            else: n += 1
    if b and not n: return 'B'
    if i and not n: return 'I'
    return 'N'

# ---------------------------------------------------------------- asset types
ASSET_CLASSES = []
for r in T[2].rows[1:]:
    ASSET_CLASSES.append(dict(asset_class=cell(r.cells[0]), description=cell(r.cells[1]),
                              geometry=cell(r.cells[2])))

TYPES = []
for r in T[3].rows[1:]:
    num = cell(r.cells[1])
    cls = cell(r.cells[0]); name = cell(r.cells[2]); desc = cell(r.cells[3])
    if not re.fullmatch(r'\d{3}', num):
        log('ASSET TYPE: non 3-digit # %r for %r' % (num, name))
        tcode = ''
    else:
        tcode = num[1:]                       # 1-digit class + 2-digit type (per Standardized names section)
    TYPES.append(OrderedDict(type_code=tcode, asset_class=cls, type_id=snake(name),
                             type_name=name, description=desc, allowed_parent_classes='',
                             _num=num))

TYPE_BY_NAME = {t['type_name']: t for t in TYPES}

# --------------------------------------------------------- attribute tables
# table index -> asset type name (from the document's own headings 6.5.x)
TYPE_TABLES = {5:'Water system',6:'Water facility',7:'Source',8:'Pump',9:'Tank',10:'Power',
    11:'Treatment',12:'Meter',13:'Electrical',14:'Valve',15:'Hydrant',16:'Junction',
    17:'Sampling point',18:'Sensor',19:'Analyzer',20:'Structure',21:'Water point',
    22:'Other vertical',23:'Pipe'}

ATTRS, CHOICES, GROUPS = [], [], []
CODE_SEEN = defaultdict(list)
INHERITED = []

def parse_table(ti, scope, atype):
    cls = TYPE_BY_NAME[atype]['asset_class'] if atype else ''
    group = ''; group_cond = ''; group_desc = ''
    cur = None
    for ri, r in enumerate(T[ti].rows):
        if ri == 0: continue
        code = cell(r.cells[0]); name = cell(r.cells[1]); desc = cell(r.cells[2])
        dt = cell(r.cells[3]); cond = cell(r.cells[4])
        k = kind(r)
        if not any([code, name, desc, dt, cond]):
            log('TABLE %d row %d: completely blank row, skipped' % (ti, ri)); continue
        if code:                                              # attribute
            if k == 'B': log('TABLE %d row %d: row has a code (%s) AND bold name %r' % (ti, ri, code, name))
            if not re.fullmatch(r'\d{5}', code):
                log('TABLE %d row %d: code %r is not 5 digits' % (ti, ri, code))
            dtype, uq = dt, ''
            m = re.match(r'^Unit:\s*(.+)$', dt)
            if m:
                dtype = 'Unit'; uq = re.sub(r'-\s+', '', m.group(1)).strip()
                if uq != m.group(1).strip():
                    log('TABLE %d row %d code %s: unit quantity %r de-hyphenated to %r (line-break artifact)'
                        % (ti, ri, code, m.group(1), uq))
            app = cond
            if not app and group_cond:
                app = group_cond
                INHERITED.append((code, name, group, group_cond))
            a = OrderedDict(code=code, name=name, scope=scope, asset_class=cls,
                            asset_type=atype or '', data_type=dtype, unit_quantity=uq,
                            description=desc,
                            applicability=app,
                            required='TRUE' if code == '00003' else '',
                            attribute_group=group)
            ATTRS.append(a); CODE_SEEN[code].append((ti, ri, name)); cur = a
        elif k == 'B':                                        # group heading
            group = name; group_cond = cond; group_desc = desc; cur = None
            GROUPS.append(OrderedDict(table=ti, asset_type=atype or '(general)', group_name=name,
                                      group_description=desc, group_condition=cond))
            if dt: log('TABLE %d row %d: group %r unexpectedly has a data type %r' % (ti, ri, name, dt))
        else:                                                 # choice option
            if cur is None:
                log('TABLE %d row %d: choice-looking row %r with no preceding attribute' % (ti, ri, name)); continue
            if cur['data_type'] != 'Choice':
                log('TABLE %d row %d: choice %r attached to attribute %s %r whose data type is %r'
                    % (ti, ri, name, cur['code'], cur['name'], cur['data_type']))
            if k != 'I':
                log('TABLE %d row %d: choice %r under %s is NOT italicised (formatting inconsistency)'
                    % (ti, ri, name, cur['code']))
            if not desc: log('TABLE %d row %d: choice %r under %s has an empty description' % (ti, ri, name, cur['code']))
            CHOICES.append(OrderedDict(attribute_code=cur['code'], choice_id=snake(name),
                                       choice_name=name, description=desc))

parse_table(4, 'general', None)
for ti in sorted(TYPE_TABLES): parse_table(ti, 'type_specific', TYPE_TABLES[ti])

# ------------------------------------------------------------------- units
UNITS = []
q = base = None
for r in T[1].rows[1:]:
    c = [cell(x) for x in r.cells]
    if len(set(c)) == 1:                     # merged quantity header row
        q = c[0]; base = None; continue
    sym, nm, fac, sysname = c[0], c[1], c[2], c[3]
    if base is None: base = sym            # first unit listed is the base unit (doc: "base unit ... appears first")
    notes = 'Name: %s; System: %s' % (nm, sysname)
    if q.endswith('*'):
        notes += '; footnote %s applies (see EXTRACTION-NOTES.md)' % ('**' if q.endswith('**') else '*')
    if q.startswith('Temperature'):
        notes += '; value in conversion_factor_to_si is a conversion FORMULA, not a factor'
    UNITS.append(OrderedDict(quantity=q.rstrip('*'), si_base_unit=base, alternate_unit=sym,
                             conversion_factor_to_si=fac, notes=notes))

# ------------------------------------------------------------------- write
def wr(fn, header, rows):
    with open(fn, 'w', newline='\n', encoding='utf-8') as f:
        w = csv.DictWriter(f, fieldnames=header, lineterminator='\n',
                           quoting=csv.QUOTE_MINIMAL, extrasaction='ignore')
        w.writeheader()
        for r in rows: w.writerow(r)

wr('asset-types.csv', ['type_code','asset_class','type_id','type_name','description','allowed_parent_classes'], TYPES)
wr('attributes.csv', ['code','name','scope','asset_class','asset_type','data_type','unit_quantity',
                      'description','applicability','required','attribute_group'], ATTRS)
wr('choices.csv', ['attribute_code','choice_id','choice_name','description'], CHOICES)
wr('units.csv', ['quantity','si_base_unit','alternate_unit','conversion_factor_to_si','notes'], UNITS)

json.dump(dict(groups=GROUPS, inherited=INHERITED, log=LOG,
               dup=[k for k,v in CODE_SEEN.items() if len(v)>1],
               dup_detail={k:v for k,v in CODE_SEEN.items() if len(v)>1},
               classes=ASSET_CLASSES,
               dtypes=sorted(set(a['data_type'] for a in ATTRS)),
               uqs=sorted(set(a['unit_quantity'] for a in ATTRS if a['unit_quantity'])),
               ), open('_meta.json','w'), indent=1)
print('types',len(TYPES),'attrs',len(ATTRS),'choices',len(CHOICES),'units',len(UNITS),'groups',len(GROUPS))
print('LOG:'); [print(' ',l) for l in LOG]
